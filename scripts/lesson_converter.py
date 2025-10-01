#!/usr/bin/env python3
"""
SATify Lesson Converter
Converts between JSON lesson format and editable text/Word documents in both directions.

Usage:
    python lesson_converter.py json_to_text lesson_01.json
    python lesson_converter.py text_to_json lesson_01.txt
    python lesson_converter.py json_to_word lesson_01.json
    python lesson_converter.py word_to_json lesson_01.docx
"""

import json
import sys
import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word document features will be disabled.")
    print("Install with: pip install python-docx")


class LessonConverter:
    """Converts SATify lessons between JSON and editable text/Word formats."""

    def __init__(self):
        self.supported_slide_types = [
            'introduction', 'concept_teaching', 'strategy_teaching',
            'guided_example', 'independent_practice', 'concept_reinforcement',
            'mastery_check'
        ]

    def json_to_text(self, json_file: str, output_file: Optional[str] = None) -> str:
        """Convert JSON lesson to editable text format."""
        print(f"Converting {json_file} to text format...")

        # Load JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            lesson_data = json.load(f)

        # Generate text content
        text_content = self._generate_text_content(lesson_data)

        # Determine output file
        if output_file is None:
            output_file = json_file.replace('.json', '.txt')

        # Write text file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text_content)

        print(f"✅ Text file created: {output_file}")
        return output_file

    def json_to_word(self, json_file: str, output_file: Optional[str] = None) -> str:
        """Convert JSON lesson to editable Word document."""
        if not DOCX_AVAILABLE:
            print("❌ Word conversion not available. Install python-docx package.")
            return None

        print(f"Converting {json_file} to Word document...")

        # Load JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            lesson_data = json.load(f)

        # Create Word document
        doc = Document()
        self._generate_word_content(lesson_data, doc)

        # Determine output file
        if output_file is None:
            output_file = json_file.replace('.json', '.docx')

        # Save Word document
        doc.save(output_file)

        print(f"✅ Word document created: {output_file}")
        return output_file

    def text_to_json(self, text_file: str, output_file: Optional[str] = None) -> str:
        """Convert text format back to JSON lesson."""
        print(f"Converting {text_file} to JSON format...")

        # Read text file
        with open(text_file, 'r', encoding='utf-8') as f:
            text_content = f.read()

        # Parse text content
        lesson_data = self._parse_text_content(text_content)

        # Determine output file
        if output_file is None:
            output_file = text_file.replace('.txt', '.json')

        # Write JSON file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(lesson_data, f, indent=2, ensure_ascii=False)

        print(f"✅ JSON file created: {output_file}")
        return output_file

    def word_to_json(self, word_file: str, output_file: Optional[str] = None) -> str:
        """Convert Word document back to JSON lesson."""
        if not DOCX_AVAILABLE:
            print("❌ Word conversion not available. Install python-docx package.")
            return None

        print(f"Converting {word_file} to JSON format...")

        # Read Word document
        doc = Document(word_file)
        text_content = self._extract_text_from_word(doc)

        # Parse content (same as text format)
        lesson_data = self._parse_text_content(text_content)

        # Determine output file
        if output_file is None:
            output_file = word_file.replace('.docx', '.json')

        # Write JSON file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(lesson_data, f, indent=2, ensure_ascii=False)

        print(f"✅ JSON file created: {output_file}")
        return output_file

    def _generate_text_content(self, lesson_data: Dict[str, Any]) -> str:
        """Generate human-readable text content from JSON lesson data."""
        lines = []

        # Header
        lines.append("=" * 80)
        lines.append("SATIFY LESSON EDITOR")
        lines.append("=" * 80)
        lines.append("")

        # Basic info
        lines.append("# LESSON INFORMATION")
        lines.append(f"ID: {lesson_data.get('id', '')}")
        lines.append(f"Title: {lesson_data.get('title', '')}")
        lines.append(f"Subtitle: {lesson_data.get('subtitle', '')}")
        lines.append(f"Level: {lesson_data.get('level', '')}")
        lines.append(f"Duration: {lesson_data.get('duration', '')}")
        lines.append(f"Skill Codes: {', '.join(lesson_data.get('skill_codes', []))}")
        lines.append("")

        # Learning objectives
        lines.append("# LEARNING OBJECTIVES")
        for i, obj in enumerate(lesson_data.get('learning_objectives', []), 1):
            lines.append(f"{i}. {obj}")
        lines.append("")

        # Success criteria
        success = lesson_data.get('success_criteria', {})
        lines.append("# SUCCESS CRITERIA")
        lines.append(f"Mastery Threshold: {success.get('mastery_threshold', 0.75)}")
        lines.append(f"Minimum Accuracy: {success.get('min_accuracy', 0.7)}")
        lines.append(f"Required Slides: {success.get('required_slides', 'all')}")
        lines.append("")

        # Slides
        lines.append("# SLIDES")
        lines.append("")

        for slide in lesson_data.get('slides', []):
            lines.extend(self._format_slide_as_text(slide))
            lines.append("")

        # Footer
        lines.append("=" * 80)
        lines.append("END OF LESSON")
        lines.append("=" * 80)
        lines.append("")
        lines.append("EDITING INSTRUCTIONS:")
        lines.append("- Edit any text content directly")
        lines.append("- Keep the structure markers (##, ---, etc.)")
        lines.append("- Don't change IDs unless creating a new lesson")
        lines.append("- Save and use text_to_json to convert back")

        return "\\n".join(lines)

    def _format_slide_as_text(self, slide: Dict[str, Any]) -> List[str]:
        """Format a single slide as editable text."""
        lines = []

        lines.append(f"## SLIDE: {slide.get('id', 'unknown')}")
        lines.append(f"Type: {slide.get('type', '')}")
        lines.append(f"Title: {slide.get('title', '')}")
        lines.append(f"Duration: {slide.get('duration_estimate', 0)} seconds")
        lines.append("")

        # Content
        content = slide.get('content', {})
        lines.append("### CONTENT")

        if 'heading' in content:
            lines.append(f"Heading: {content['heading']}")

        if 'text' in content:
            lines.append("Text:")
            lines.append(f"    {content['text']}")

        if 'bullet_points' in content:
            lines.append("Bullet Points:")
            for point in content['bullet_points']:
                lines.append(f"    • {point}")

        if 'strategy_steps' in content:
            lines.append("Strategy Steps:")
            for step in content['strategy_steps']:
                lines.append(f"    Step {step.get('step', 1)}: {step.get('title', '')}")
                lines.append(f"        {step.get('description', '')}")
                if 'example' in step:
                    lines.append(f"        Example: {step['example']}")

        if 'concept_box' in content:
            box = content['concept_box']
            lines.append(f"Concept Box: {box.get('title', '')}")
            for point in box.get('points', []):
                lines.append(f"    • {point}")

        if 'examples' in content:
            examples = content['examples']
            lines.append("Examples:")
            for key, value in examples.items():
                lines.append(f"    {key}: {value}")

        if 'practice_transition' in content:
            pt = content['practice_transition']
            lines.append("Practice Transition:")
            lines.append(f"    Text: {pt.get('text', '')}")
            lines.append(f"    Button: {pt.get('button_text', '')}")

        # Interactions
        if 'interactions' in slide:
            lines.append("### INTERACTIONS")
            for interaction in slide['interactions']:
                lines.append(f"Type: {interaction.get('type', '')}")
                if 'text' in interaction:
                    lines.append(f"Text: {interaction['text']}")
                if 'elements' in interaction:
                    for elem in interaction['elements']:
                        lines.append(f"    Trigger: {elem.get('trigger', '')}")
                        lines.append(f"    Reveal: {elem.get('reveal', '')}")

        lines.append("---")  # Slide separator

        return lines

    def _generate_word_content(self, lesson_data: Dict[str, Any], doc: Document):
        """Generate Word document content from JSON lesson data."""
        # Title
        title = doc.add_heading(lesson_data.get('title', 'Untitled Lesson'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if lesson_data.get('subtitle'):
            subtitle = doc.add_heading(lesson_data['subtitle'], level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Basic info table
        doc.add_heading('Lesson Information', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        info_rows = [
            ('ID', lesson_data.get('id', '')),
            ('Level', lesson_data.get('level', '')),
            ('Duration', lesson_data.get('duration', '')),
            ('Skill Codes', ', '.join(lesson_data.get('skill_codes', []))),
        ]

        for i, (label, value) in enumerate(info_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)

        # Learning objectives
        doc.add_heading('Learning Objectives', level=1)
        for obj in lesson_data.get('learning_objectives', []):
            p = doc.add_paragraph(obj, style='List Bullet')

        # Success criteria
        success = lesson_data.get('success_criteria', {})
        doc.add_heading('Success Criteria', level=1)
        doc.add_paragraph(f"Mastery Threshold: {success.get('mastery_threshold', 0.75)}")
        doc.add_paragraph(f"Minimum Accuracy: {success.get('min_accuracy', 0.7)}")
        doc.add_paragraph(f"Required Slides: {success.get('required_slides', 'all')}")

        # Slides
        doc.add_page_break()
        doc.add_heading('Lesson Slides', level=1)

        for i, slide in enumerate(lesson_data.get('slides', []), 1):
            self._format_slide_as_word(slide, doc, i)

    def _format_slide_as_word(self, slide: Dict[str, Any], doc: Document, slide_num: int):
        """Format a single slide in Word document."""
        # Slide header
        doc.add_heading(f"Slide {slide_num}: {slide.get('title', 'Untitled')}", level=2)

        # Slide info
        doc.add_paragraph(f"ID: {slide.get('id', '')}")
        doc.add_paragraph(f"Type: {slide.get('type', '')}")
        doc.add_paragraph(f"Duration: {slide.get('duration_estimate', 0)} seconds")

        # Content
        content = slide.get('content', {})

        if content.get('heading'):
            doc.add_heading(content['heading'], level=3)

        if content.get('text'):
            doc.add_paragraph(content['text'])

        if content.get('bullet_points'):
            doc.add_heading('Key Points', level=4)
            for point in content['bullet_points']:
                doc.add_paragraph(point, style='List Bullet')

        if content.get('strategy_steps'):
            doc.add_heading('Strategy Steps', level=4)
            for step in content['strategy_steps']:
                p = doc.add_paragraph()
                p.add_run(f"Step {step.get('step', 1)}: {step.get('title', '')}").bold = True
                doc.add_paragraph(step.get('description', ''))
                if step.get('example'):
                    doc.add_paragraph(f"Example: {step['example']}", style='Intense Quote')

        if content.get('concept_box'):
            box = content['concept_box']
            doc.add_heading(f"Concept: {box.get('title', '')}", level=4)
            for point in box.get('points', []):
                doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph()  # Spacing

    def _parse_text_content(self, text_content: str) -> Dict[str, Any]:
        """Parse editable text content back to JSON lesson data."""
        # First, try to detect if this is the standard SATify format
        if '# LESSON INFORMATION' in text_content or '## SLIDE:' in text_content:
            return self._parse_standard_format(text_content)
        else:
            # Try to parse as custom format
            return self._parse_custom_format(text_content)

    def _parse_standard_format(self, text_content: str) -> Dict[str, Any]:
        """Parse standard SATify lesson text format."""
        lines = text_content.split('\\n')
        lesson_data = {}
        current_section = None
        current_slide = None
        slides = []

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith('# LESSON INFORMATION'):
                current_section = 'info'
                i += 1
                continue

            elif line.startswith('# LEARNING OBJECTIVES'):
                current_section = 'objectives'
                lesson_data['learning_objectives'] = []
                i += 1
                continue

            elif line.startswith('# SUCCESS CRITERIA'):
                current_section = 'success'
                lesson_data['success_criteria'] = {}
                i += 1
                continue

            elif line.startswith('# SLIDES'):
                current_section = 'slides'
                i += 1
                continue

            elif line.startswith('## SLIDE:'):
                if current_slide:
                    slides.append(current_slide)
                current_slide = self._parse_slide_header(lines, i)
                i += 1
                continue

            elif line.startswith('---') and current_slide:
                slides.append(current_slide)
                current_slide = None
                i += 1
                continue

            # Parse content based on current section
            if current_section == 'info' and ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().lower().replace(' ', '_')
                value = value.strip()

                if key == 'skill_codes':
                    lesson_data[key] = [code.strip() for code in value.split(',')]
                else:
                    lesson_data[key] = value

            elif current_section == 'objectives' and line and line[0].isdigit():
                # Remove number prefix
                obj = re.sub(r'^\\d+\\.\\s*', '', line)
                lesson_data['learning_objectives'].append(obj)

            elif current_section == 'success' and ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().lower().replace(' ', '_')
                value = value.strip()

                try:
                    if key in ['mastery_threshold', 'min_accuracy']:
                        value = float(value)
                except ValueError:
                    pass

                lesson_data['success_criteria'][key] = value

            elif current_slide:
                current_slide = self._parse_slide_content(current_slide, lines, i)

            i += 1

        # Add final slide if exists
        if current_slide:
            slides.append(current_slide)

        lesson_data['slides'] = slides
        return lesson_data

    def _parse_custom_format(self, text_content: str) -> Dict[str, Any]:
        """Parse custom lesson format like the one in prabha directory."""
        lines = [line.strip() for line in text_content.split('\\n') if line.strip()]
        lesson_data = {
            'slides': []
        }

        slides = []
        current_slide = None
        slide_counter = 1

        i = 0
        while i < len(lines):
            line = lines[i]

            # Check for lesson title (first meaningful line)
            if i == 0 and not lesson_data.get('title'):
                lesson_data['title'] = line
                lesson_data['id'] = 'lesson_custom'
                lesson_data['skill_codes'] = []

            # Look for lesson information section
            elif 'Lesson Information' in line:
                i += 1
                continue
            elif line.startswith('Strategy-First:'):
                lesson_data['subtitle'] = line.replace('Strategy-First:', '').strip()
            elif 'Your Learning Objective:' in line or 'The Core:' in line:
                if not lesson_data.get('learning_objectives'):
                    lesson_data['learning_objectives'] = []
                if i + 1 < len(lines):
                    lesson_data['learning_objectives'].append(lines[i + 1])
            elif 'Mastery Threshold:' in line:
                if not lesson_data.get('success_criteria'):
                    lesson_data['success_criteria'] = {}
                lesson_data['success_criteria']['mastery_threshold'] = float(line.split(':')[1].strip())
            elif 'Minimum Accuracy:' in line:
                if not lesson_data.get('success_criteria'):
                    lesson_data['success_criteria'] = {}
                lesson_data['success_criteria']['min_accuracy'] = float(line.split(':')[1].strip())

            # Look for slide markers
            elif line.startswith('Slide ') and ':' in line:
                # Save previous slide
                if current_slide:
                    slides.append(current_slide)

                # Create new slide
                slide_title = line.split(':', 1)[1].strip()
                current_slide = {
                    'id': f'slide_{slide_counter:02d}',
                    'type': 'concept_teaching',  # Default type
                    'title': slide_title,
                    'duration_estimate': 240,  # Default duration
                    'content': {
                        'heading': slide_title,
                        'text': '',
                        'bullet_points': []
                    }
                }
                slide_counter += 1

            # Look for slide metadata
            elif current_slide and line.startswith('ID:'):
                current_slide['id'] = line.split(':', 1)[1].strip()
            elif current_slide and line.startswith('Type:'):
                slide_type = line.split(':', 1)[1].strip().lower().replace(' ', '_')
                current_slide['type'] = slide_type
            elif current_slide and line.startswith('Duration:'):
                duration_str = line.split(':', 1)[1].strip()
                try:
                    duration = int(duration_str.split()[0])
                    current_slide['duration_estimate'] = duration
                except:
                    pass

            # Collect slide content
            elif current_slide:
                if line.startswith('*') or line.startswith('•'):
                    # Bullet point
                    bullet = line.lstrip('*• ').strip()
                    if bullet:
                        current_slide['content']['bullet_points'].append(bullet)
                elif line and not line.startswith('Slide '):
                    # Regular content text
                    if current_slide['content']['text']:
                        current_slide['content']['text'] += '\\n' + line
                    else:
                        current_slide['content']['text'] = line

            i += 1

        # Don't forget the last slide
        if current_slide:
            slides.append(current_slide)

        lesson_data['slides'] = slides

        # Set some defaults if not found
        if not lesson_data.get('level'):
            lesson_data['level'] = 'Foundation'
        if not lesson_data.get('duration'):
            lesson_data['duration'] = f'{len(slides) * 4}-{len(slides) * 6} min'

        return lesson_data

    def _parse_slide_header(self, lines: List[str], start_idx: int) -> Dict[str, Any]:
        """Parse slide header information."""
        slide = {}

        # Extract slide ID from header
        header_line = lines[start_idx].strip()
        slide_id = header_line.split(':', 1)[1].strip()
        slide['id'] = slide_id

        # Parse slide metadata
        for i in range(start_idx + 1, len(lines)):
            line = lines[i].strip()

            if line.startswith('###') or line.startswith('---') or not line:
                break

            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().lower()
                value = value.strip()

                if key == 'duration':
                    try:
                        slide['duration_estimate'] = int(value.split()[0])
                    except (ValueError, IndexError):
                        slide['duration_estimate'] = 0
                else:
                    slide[key] = value

        slide['content'] = {}
        slide['interactions'] = []

        return slide

    def _parse_slide_content(self, slide: Dict[str, Any], lines: List[str], line_idx: int) -> Dict[str, Any]:
        """Parse slide content from text lines."""
        line = lines[line_idx].strip()

        # Skip if not content line
        if not line or line.startswith('#') or line.startswith('---'):
            return slide

        content = slide.get('content', {})

        # Parse different content types
        if line.startswith('Heading:'):
            content['heading'] = line.split(':', 1)[1].strip()

        elif line.startswith('Text:'):
            # Look for indented text on next line
            if line_idx + 1 < len(lines) and lines[line_idx + 1].startswith('    '):
                content['text'] = lines[line_idx + 1].strip()

        elif line.startswith('Bullet Points:'):
            content['bullet_points'] = self._parse_bullet_points(lines, line_idx + 1)

        elif line.startswith('Strategy Steps:'):
            content['strategy_steps'] = self._parse_strategy_steps(lines, line_idx + 1)

        slide['content'] = content
        return slide

    def _parse_bullet_points(self, lines: List[str], start_idx: int) -> List[str]:
        """Parse bullet points from indented lines."""
        points = []

        for i in range(start_idx, len(lines)):
            line = lines[i].strip()
            if not line or not line.startswith('•'):
                break

            point = line[1:].strip()  # Remove bullet
            points.append(point)

        return points

    def _parse_strategy_steps(self, lines: List[str], start_idx: int) -> List[Dict[str, Any]]:
        """Parse strategy steps from text."""
        steps = []
        current_step = None

        for i in range(start_idx, len(lines)):
            line = lines[i].strip()

            if not line or line.startswith('###'):
                break

            if line.startswith('Step '):
                if current_step:
                    steps.append(current_step)

                # Parse step header
                parts = line.split(':', 1)
                if len(parts) == 2:
                    step_part = parts[0].replace('Step ', '').strip()
                    try:
                        step_num = int(step_part)
                    except ValueError:
                        step_num = len(steps) + 1

                    current_step = {
                        'step': step_num,
                        'title': parts[1].strip(),
                        'description': ''
                    }

            elif current_step and line.startswith('        '):
                # Continuation of step description or example
                clean_line = line.strip()
                if clean_line.startswith('Example:'):
                    current_step['example'] = clean_line[8:].strip()
                else:
                    current_step['description'] = clean_line

        if current_step:
            steps.append(current_step)

        return steps

    def _extract_text_from_word(self, doc: Document) -> str:
        """Extract text content from Word document."""
        text_lines = []

        for paragraph in doc.paragraphs:
            text_lines.append(paragraph.text)

        return '\\n'.join(text_lines)


def main():
    """Command-line interface for lesson converter."""
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python lesson_converter.py json_to_text <input.json> [output.txt]")
        print("  python lesson_converter.py json_to_word <input.json> [output.docx]")
        print("  python lesson_converter.py text_to_json <input.txt> [output.json]")
        print("  python lesson_converter.py word_to_json <input.docx> [output.json]")
        return

    action = sys.argv[1]
    input_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else None

    converter = LessonConverter()

    try:
        if action == 'json_to_text':
            converter.json_to_text(input_file, output_file)
        elif action == 'json_to_word':
            converter.json_to_word(input_file, output_file)
        elif action == 'text_to_json':
            converter.text_to_json(input_file, output_file)
        elif action == 'word_to_json':
            converter.word_to_json(input_file, output_file)
        else:
            print(f"Unknown action: {action}")
            print("Valid actions: json_to_text, json_to_word, text_to_json, word_to_json")

    except FileNotFoundError:
        print(f"❌ File not found: {input_file}")
    except Exception as e:
        print(f"❌ Error: {e}")


if __name__ == '__main__':
    main()